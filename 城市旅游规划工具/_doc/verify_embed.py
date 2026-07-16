from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

root = Path(__file__).resolve().parents[1]
report = next(root.glob("*FDE*.docx"))
charts = root / "_doc" / "charts"
target = root / "_doc" / "embed_verify.docx"
figures = [
    (3, "fig_3-1_system_three_layer_architecture.png"),
    (5, "fig_3-2_function_module_relationship.png"),
    (6, "fig_3-3_user_llm_tool_interaction_flow.png"),
]

doc = Document(report)
for table_index, filename in figures:
    cell = doc.tables[table_index].cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(charts / filename), width=Cm(14.7))
doc.save(target)

check = Document(target)
with ZipFile(target) as archive:
    media = [name for name in archive.namelist() if name.startswith("word/media/") and name.endswith(".png")]
print("TARGET", target)
print("INLINE", len(check.inline_shapes))
print("MEDIA", len(media), media)
for index, _ in figures:
    print("TABLE", index, "DRAWINGS", len(check.tables[index].cell(0, 0)._tc.xpath(".//w:drawing")))
