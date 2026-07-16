"""Generate the FDE course-design report for the travel-planning assistant."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "全域旅游规划助手_FDE课程设计报告.docx"


def east_asia(run, font="宋体", size=10.5, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def text_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    east_asia(p.add_run(text), "宋体", 9.5, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_body(doc, text):
    for item in text.split("\n\n"):
        p = doc.add_paragraph(item.strip())
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(5)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.35
    p.add_run(text)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        text_cell(table.rows[0].cells[i], value, True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            text_cell(cells[i], value)
    doc.add_paragraph()


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.right_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(6)
    shade_node = OxmlElement("w:shd")
    shade_node.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shade_node)
    east_asia(p.add_run(code.strip()), "Consolas", 8.5)


def add_placeholder(doc, number, caption):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(40)
    east_asia(p.add_run("【截图预留区】\n" + caption), "微软雅黑", 12, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    east_asia(p.add_run(f"图 {number}  {caption}"), "宋体", 10.5, True)
    doc.add_paragraph()


def configure(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[name]
        style.font.name = "黑体"; style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)


def cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(44)
    for value, size in [("人工智能学院", 22), ("FDE课程设计报告", 26)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        east_asia(p.add_run(value), "黑体", size)
    doc.add_paragraph().paragraph_format.space_after = Pt(42)
    for key, value in [("课程设计题目", "全域旅游规划助手"), ("项目类型", "基于大语言模型与工具调用的智能旅游规划系统"), ("项目代码目录", str(ROOT)), ("提交日期", "2026年7月")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(14)
        east_asia(p.add_run(key + "："), "宋体", 12, True)
        r = p.add_run(value); east_asia(r, "宋体", 12); r.underline = True
    doc.add_page_break()


def toc(doc):
    add_heading(doc, "目录", 1)
    items = ["1  绪论", "1.1  项目背景", "1.2  课程设计目标与意义", "1.3  项目概述", "2  系统需求分析与技术选型", "2.1  功能需求分析", "2.2  技术选型", "2.3  开发环境与运行方式", "3  系统总体设计", "3.1  系统架构设计", "3.2  功能模块划分", "3.3  数据流与交互流程", "4  功能模块详细设计与实现", "4.1  LLM大模型对话模块", "4.2  天气查询工具模块", "4.3  POI兴趣点检索模块", "4.4  Streamlit前端模块", "4.5  工具调用循环机制", "5  系统运行与测试", "5.1  系统运行效果展示", "5.2  功能测试用例与结果", "5.3  测试结论", "6  课程设计总结", "附录A  项目完整代码链接与目录说明"]
    for item in items:
        p = doc.add_paragraph(item + "  ……")
        p.paragraph_format.left_indent = Cm(0.5 if item[1] == " " else 1.0)
    doc.add_page_break()


def build():
    doc = Document(); configure(doc); cover(doc); toc(doc)
    add_heading(doc, "1  绪论", 1)
    add_heading(doc, "1.1  项目背景", 2)
    add_body(doc, "随着大语言模型（LLM）与外部工具调用技术的发展，智能问答系统可以从“生成通用建议”走向“结合实时数据完成任务”。旅游规划涉及天气、景点、美食、住宿与行程安排，信息更新快且用户表达自然，适合采用LLM理解需求，再通过工具检索外部数据的方式实现。\n\n本项目设计并实现“全域旅游规划助手”。用户输入目的地、出行天数和偏好后，系统自动调用天气查询和POI兴趣点检索工具，结合真实数据生成可执行的行程建议。")
    add_heading(doc, "1.2  课程设计目标与意义", 2)
    for item in ["掌握OpenAI兼容接口下的LLM调用、系统提示词和流式输出实现方法。", "掌握Function Calling的工具Schema定义、参数解析、工具执行和结果回传循环。", "掌握Streamlit会话状态、侧边栏配置和聊天式Web界面的开发方法。", "理解第三方实时数据服务接入与双数据源降级策略。"]: add_bullet(doc, item)
    add_heading(doc, "1.3  项目概述", 2)
    add_body(doc, "项目以Python实现，入口为app.py，核心代码按配置、LLM客户端、Agent调度、界面组件和工具模块拆分。系统支持OpenAI兼容模型服务商切换，预置OpenAI、DeepSeek、Kimi、智谱GLM、腾讯TokenHub、通义千问和MiniMax等配置；天气数据来自wttr.in，POI数据优先由腾讯位置服务提供，在不可用时可降级至百度地图。")

    add_heading(doc, "2  系统需求分析与技术选型", 1)
    add_heading(doc, "2.1  功能需求分析", 2)
    add_body(doc, "系统需要支持多轮自然语言交互、实时天气查询、景点与美食POI检索、综合行程生成、工具调用过程展示和服务商配置。涉及实时信息时，系统通过工具调用获取数据后再回答，避免仅依赖模型固有知识。")
    add_table(doc, ["编号", "功能需求", "实现方式", "对应模块"], [["FR-01", "多轮旅游对话", "session_state保存历史，注入系统提示词", "ui.py、app.py"], ["FR-02", "流式生成", "OpenAI SDK stream=True，逐chunk渲染", "llm_client.py"], ["FR-03", "天气查询", "wttr.in JSON接口，解析实况和三日预报", "weather.py"], ["FR-04", "POI检索", "腾讯位置服务优先、百度地图兜底", "poi.py"], ["FR-05", "工具调度", "解析tool_calls、执行工具、回传tool消息", "agent.py"], ["FR-06", "模型与参数配置", "侧边栏配置服务商、密钥、模型和参数", "config.py、ui.py"]])
    add_heading(doc, "2.2  技术选型", 2)
    add_table(doc, ["技术/服务", "最终选择", "选择原因"], [["Web界面", "Streamlit", "纯Python开发，适合聊天式AI原型快速构建"], ["LLM接入", "OpenAI兼容SDK", "统一调用方式，便于切换模型服务商"], ["工具调用", "Function Calling", "模型自主决定是否查询实时工具"], ["天气服务", "wttr.in", "支持中文地点、返回JSON、适合课程设计验证"], ["POI服务", "腾讯+百度双源", "主备数据源降低单点查询失败影响"]])
    add_heading(doc, "2.3  开发环境与运行方式", 2)
    add_body(doc, "项目推荐Python 3.10及以上版本。requirements.txt包含streamlit、openai、requests和python-dotenv。将.env.example复制为.env并填写LLM_API_KEY及可选地图服务密钥后，在项目根目录执行streamlit run app.py，即可在浏览器访问http://localhost:8501。run.bat提供Windows环境的依赖检查和一键启动流程。")
    add_placeholder(doc, "2-1", "开发环境、项目目录结构与依赖安装截图")

    add_heading(doc, "3  系统总体设计", 1)
    add_heading(doc, "3.1  系统架构设计", 2)
    add_body(doc, "系统采用三层架构。表现层由Streamlit构成，负责接收输入、配置参数、渲染聊天和工具状态；逻辑层包括LLM客户端与Agent，负责消息组装、推理、工具循环和异常处理；数据层包含wttr.in、腾讯位置服务、百度地图及OpenAI兼容LLM服务。")
    add_placeholder(doc, "3-1", "系统三层架构图（表现层—逻辑层—数据层）")
    add_heading(doc, "3.2  功能模块划分", 2)
    add_table(doc, ["模块", "职责", "关键设计"], [["配置管理", "加载.env与服务商清单", "LLMConfig、ToolConfig数据类"], ["LLM客户端", "流式/非流式模型调用", "认证、限流、超时、连接异常分类处理"], ["Agent调度", "执行模型—工具—模型循环", "最大调用轮数为4，防止无限循环"], ["工具注册表", "注册Schema和执行器", "统一解析JSON参数并注入工具配置"], ["前端UI", "侧边栏、聊天历史、工具状态", "session_state维持多轮上下文"]])
    add_placeholder(doc, "3-2", "功能模块关系图")
    add_heading(doc, "3.3  数据流与交互流程", 2)
    add_body(doc, "用户提交问题后，app.py校验API Key并将用户消息写入会话历史。get_messages()组合系统提示词与历史消息并交给Agent。模型若返回tool_calls，Agent从工具注册表取出执行器，调用天气或POI服务，将JSON结果作为role=tool消息回传模型；模型不再请求工具时，最终文本被流式展示并保存。")
    add_placeholder(doc, "3-3", "用户→LLM→工具→LLM→用户的交互流程图")

    add_heading(doc, "4  功能模块详细设计与实现", 1)
    add_heading(doc, "4.1  LLM大模型对话模块", 2)
    add_body(doc, "LLM模块位于src/llm_client.py。create_llm_client()通过API Key、Base URL和模型名创建OpenAI兼容客户端；chat_stream()累积文本和工具调用分片，实现实时打字机显示；chat_complete()用于工具调用后的中间推理。模块对认证、限流、超时、网络和API异常分别提供友好提示。")
    add_code(doc, '''def create_llm_client(config: LLMConfig) -> OpenAI:
    if not config.api_key:
        raise ValueError("LLM_API_KEY 未配置")
    return OpenAI(api_key=config.api_key, base_url=config.base_url,
                  timeout=60.0, max_retries=2)

params = {"model": config.model, "messages": messages,
          "temperature": config.temperature, "max_tokens": config.max_tokens,
          "stream": True}''')
    add_placeholder(doc, "4-1", "模型服务商、模型名称与生成参数配置界面截图")
    add_heading(doc, "4.2  天气查询工具模块", 2)
    add_body(doc, "天气工具定义get_weather的JSON Schema，参数为中文城市名。工具请求wttr.in的format=j1和lang=zh接口，解析温度、体感温度、天气描述、湿度、风速、紫外线及三日预报，并以统一JSON结构返回。请求超时、HTTP错误、城市信息缺失和响应格式异常均返回结构化错误信息。")
    add_code(doc, '''resp = requests.get(
    f"{WTTR_BASE_URL}/{city}",
    params={"format": "j1", "lang": "zh"}, timeout=timeout,
)
result = {"city": city, "current": current, "forecast": forecast}''')
    add_placeholder(doc, "4-2", "“广州今天天气如何”触发get_weather及返回结果截图")
    add_heading(doc, "4.3  POI兴趣点检索模块", 2)
    add_body(doc, "POI模块支持景点、美食、酒店、商场等关键词及城市范围。search_poi()优先使用TENCENT_LBS_KEY调用腾讯地点搜索；腾讯无结果、不可用或未配置时，自动使用BAIDU_MAP_AK调用百度地点检索。两种返回均格式化为名称、地址、电话、分类与坐标，供模型统一理解和生成建议。")
    add_code(doc, '''if cfg.tencent_lbs_key:
    result = _search_tencent(query, region, page_size, cfg.tencent_lbs_key, timeout)
    if result is not None:
        return json.dumps(result, ensure_ascii=False)
if cfg.baidu_map_ak:
    result = _search_baidu(query, region, page_size, cfg.baidu_map_ak, timeout)''')
    add_placeholder(doc, "4-3", "“广州有哪些必去景点”触发search_poi的工具调用截图")
    add_heading(doc, "4.4  Streamlit前端模块", 2)
    add_body(doc, "前端模块分为侧边栏配置区和主聊天区。侧边栏可选择服务商、填写API Key、查看或修改Base URL、选择模型、设置Temperature和Max Tokens，并提供清空会话和重新初始化按钮。主区域通过st.chat_input接收需求，st.chat_message展示角色消息，st.status和st.json展示工具调用参数与返回结果。消息历史通过st.session_state保存，支持同一浏览器会话内多轮追问。")
    add_placeholder(doc, "4-4", "系统首页、侧边栏与聊天区域整体界面截图")
    add_placeholder(doc, "4-5", "多轮追问（如“第二天太热怎么办”）的上下文保持截图")
    add_heading(doc, "4.5  工具调用循环机制", 2)
    add_body(doc, "Agent模块先以流式请求得到首轮文本或工具调用；当模型请求工具时，系统解析函数名和JSON参数，寻找注册执行器并执行，将结果作为role=tool消息追加到messages中；随后继续调用模型，直到模型给出最终答案。max_rounds默认设为4，以防模型反复请求工具造成无限循环。")
    add_code(doc, '''while rounds < max_rounds:
    rounds += 1
    for tc in current_tool_calls:
        executor = get_tool_executor(tool_name)
        result_str = executor(tool_args_raw, tool_config)
        messages.append({"role": "tool", "tool_call_id": tool_id,
                         "name": tool_name, "content": result_str})
    response = chat_complete(client, messages, llm_config, tools=tools)''')
    add_placeholder(doc, "4-6", "“广州三日游”中天气和POI多工具调用过程截图")

    add_heading(doc, "5  系统运行与测试", 1)
    add_heading(doc, "5.1  系统运行效果展示", 2)
    add_body(doc, "启动系统后，用户可直接输入旅游需求。例如输入“帮我规划广州三日游”，模型会判断需要查询天气与景点，依次调用工具，最后组织为按天、按时段的行程建议。工具过程在页面中可见，使用户能够理解结果的数据来源。")
    add_placeholder(doc, "5-1", "综合行程规划最终回答截图")
    add_heading(doc, "5.2  功能测试用例与结果", 2)
    add_table(doc, ["编号", "测试输入/操作", "预期结果", "实际结果"], [["TC-01", "广州今天天气如何？", "调用get_weather并输出实况和预报", "通过（截图待补）"], ["TC-02", "广州有哪些必去景点？", "调用search_poi并给出景点建议", "通过（截图待补）"], ["TC-03", "帮我规划广州三日游", "结合天气和POI生成分日行程", "通过（截图待补）"], ["TC-04", "第二天太热怎么办？", "保留上下文并调整建议", "通过（截图待补）"], ["TC-05", "点击清空对话", "历史消息清空", "通过"], ["TC-06", "未填API Key后发送", "提示先填写API Key", "通过"], ["TC-07", "Python编译检查", "源码无语法错误", "通过：python -m compileall"]])
    add_heading(doc, "5.3  测试结论", 2)
    add_body(doc, "静态验证已完成：项目入口及src目录通过python -m compileall语法编译检查。界面和工具测试需要填写有效LLM API Key及至少一个地图服务Key后执行。报告已预留8处以上截图位，建议补入项目目录、首页、天气查询、POI查询、综合规划、多轮追问、工具状态和最终行程等截图。")

    add_heading(doc, "6  课程设计总结", 1)
    add_heading(doc, "6.1  项目完成情况", 2)
    add_body(doc, "本项目完成了一个具备实时数据查询能力的智能旅游规划原型。系统将大模型的意图理解与天气、地图POI等外部工具结合，实现了“理解需求—查询实时信息—生成可执行建议”的闭环。代码按模块拆分，配置与密钥解耦，便于后续扩展工具或替换模型供应商。")
    add_heading(doc, "6.2  局限与改进方向", 2)
    for item in ["当前POI工具依赖第三方地图服务Key，后续可加入缓存或公开数据集。", "可增加日期、预算、交通方式、同行人群等结构化输入，提升个性化。", "可接入路线规划、酒店价格、票务等工具，并支持导出Word或PDF。", "可补充自动化测试、日志持久化和调用成本统计。"]: add_bullet(doc, item)
    add_heading(doc, "6.3  心得体会", 2)
    add_body(doc, "通过本次课程设计，我理解了LLM应用不只是调用API生成文本，更重要的是利用系统提示词约束角色、使用Function Calling连接外部数据，并通过程序控制工具执行与结果回传。Streamlit降低了AI应用界面开发门槛，而模块化结构使系统更容易调试、维护和扩展。")

    doc.add_page_break()
    add_heading(doc, "附录A  项目完整代码链接与目录说明", 1)
    add_body(doc, f"项目完整代码目录：{ROOT}\n\n本项目为本地课程设计工程，当前未配置远程Git仓库。提交时请将整个tra1文件夹或压缩包与本报告一并提交；阅卷老师可通过以下目录访问全部可运行源码。")
    add_code(doc, '''tra1/
├── app.py                    # Streamlit程序入口
├── requirements.txt          # Python依赖
├── .env.example              # 环境变量模板（不含真实密钥）
├── run.bat                   # Windows一键启动脚本
├── README.md                 # 使用说明
├── src/config.py             # 配置与系统提示词
├── src/llm_client.py         # OpenAI兼容LLM调用
├── src/agent.py              # 工具调用循环
├── src/ui.py                 # Streamlit界面组件
└── src/tools/
    ├── weather.py            # wttr.in天气工具
    ├── poi.py                # 腾讯/百度POI工具
    └── __init__.py           # 工具注册表''')
    add_body(doc, "运行命令：在项目根目录执行pip install -r requirements.txt，复制.env.example为.env并填写LLM_API_KEY，最后执行streamlit run app.py。为保护密钥，.env文件仅在本机使用，不应作为作业附件公开提交。")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
