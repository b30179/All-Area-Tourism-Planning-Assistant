"""Embed three generated diagrams into their known report table placeholders."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "全域旅游规划助手_FDE课程设计报告.docx"
CHARTS = ROOT / "_doc" / "charts"

# These table indexes are stable in generate_fde_report.py: Figure 2-1 is table 2,
# then Figures 3-1, 3-2, and 3-3 are tables 3, 5, and 6.
FIGURES = [
    (3, "fig_3-1_system_three_layer_architecture.png"),
    (5, "fig_3-2_function_module_relationship.png"),
    (6, "fig_3-3_user_llm_tool_interaction_flow.png"),
]


def main():
    doc = Document(REPORT)
    for table_index, filename in FIGURES:
        cell = doc.tables[table_index].cell(0, 0)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(CHARTS / filename), width=Cm(14.7))
    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
